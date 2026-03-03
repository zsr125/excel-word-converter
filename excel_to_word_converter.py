#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Excel/Numbers 转 Word 产品技术手册工具（内容重组为分级大纲）

将 Excel (.xlsx) 或 Apple Numbers (.numbers) 工作簿转化为一份「纯文字 + 分级大纲」的 Word 文档，
不是简单格式转换，而是内容呈现方式的改变：
- 每个 sheet 名称 → Word 的一级大纲（标题 1）
- 每个 sheet 内的每一行：第一列（或第一个非空单元格）→ 二级大纲（标题 2），其余列 → 正文段落（列名：内容）
- 全文无表格，仅保留标题层级与纯文字，信息一字不差
"""

import argparse
import re
import sys
from pathlib import Path

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from numbers_parser import Document as NumbersDocument
except ImportError:
    NumbersDocument = None

try:
    from docx import Document as WordDocument
    from docx.shared import Pt, Cm, Cm
    from docx.oxml.ns import qn
except ImportError:
    WordDocument = None

# 中文数字与大纲符（一、二、三…；（一）（二）…；①②③…）
_CN_DIGIT = "零一二三四五六七八九"
def _outline_symbol_l1(i: int) -> str:
    """一级大纲符：一、二、三、…、十、十一、…"""
    n = i + 1
    if n <= 10:
        return (_CN_DIGIT[n] if n < 10 else "十") + "、"
    if n < 20:
        return "十" + _CN_DIGIT[n - 10] + "、"
    if n < 100:
        a, b = n // 10, n % 10
        return _CN_DIGIT[a] + "十" + (_CN_DIGIT[b] if b else "") + "、"
    return str(n) + "、"

def _outline_symbol_l2(i: int) -> str:
    """二级大纲符：（一）（二）（三）…"""
    n = i + 1
    if n <= 10:
        s = _CN_DIGIT[n] if n < 10 else "十"
    elif n < 20:
        s = "十" + _CN_DIGIT[n - 10]
    elif n < 100:
        a, b = n // 10, n % 10
        s = _CN_DIGIT[a] + "十" + (_CN_DIGIT[b] if b else "")
    else:
        s = str(n)
    return "（" + s + "） "

def _outline_symbol_l3(i: int) -> str:
    """三级大纲符：①②③…"""
    circles = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳"
    return (circles[i] if i < len(circles) else str(i + 1) + "．") + " "


# ---------- Excel (openpyxl) ----------

def _excel_merged_cell_value(sheet, row, col):
    """获取 Excel 单元格值，若在合并区域内则返回合并区域左上角的值。"""
    cell = sheet.cell(row=row, column=col)
    if cell.value is not None and str(cell.value).strip() != "":
        return cell.value
    for merged in sheet.merged_cells.ranges:
        if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
            return sheet.cell(row=merged.min_row, column=merged.min_col).value
    return cell.value


def _excel_sheet_to_matrix(sheet):
    """将 Excel sheet 转为二维矩阵，正确处理合并单元格。"""
    if sheet.max_row == 0 and sheet.max_column == 0:
        return []
    rows = []
    for r in range(1, sheet.max_row + 1):
        row_data = []
        for c in range(1, sheet.max_column + 1):
            val = _excel_merged_cell_value(sheet, r, c)
            if val is None:
                row_data.append("")
            else:
                row_data.append(str(val).strip() if isinstance(val, str) else str(val))
        rows.append(row_data)
    max_cols = sheet.max_column
    for row in rows:
        while len(row) < max_cols:
            row.append("")
    return rows


def _load_excel_sheets(path: Path):
    """从 Excel 文件加载所有 sheet，逐个 yield (sheet_name, [matrix])。"""
    wb = openpyxl.load_workbook(path, data_only=True)
    for sheet_name in wb.sheetnames:
        matrix = _excel_sheet_to_matrix(wb[sheet_name])
        yield sheet_name, [matrix]


# ---------- Numbers (numbers-parser) ----------

def _a1_to_rowcol(a1: str):
    """将 A1 表示法转为 (row, col)，0-based。例如 'A1' -> (0, 0), 'B3' -> (2, 1)。"""
    m = re.match(r"^([A-Z]+)(\d+)$", a1.upper())
    if not m:
        return None
    col_str, row_str = m.group(1), m.group(2)
    col = 0
    for c in col_str:
        col = col * 26 + (ord(c) - ord("A") + 1)
    return (int(row_str) - 1, col - 1)


def _numbers_cell_display_value(cell, table, row_idx: int, col_idx: int):
    """获取 Numbers 单元格的显示值，合并单元格取左上角值。"""
    if cell.value is not None:
        if hasattr(cell, "formatted_value") and cell.formatted_value is not None:
            return str(cell.formatted_value)
        return str(cell.value) if not isinstance(cell.value, str) else cell.value
    # MergedCell 或 EmptyCell：尝试从 merge_ranges 取左上角
    try:
        for range_str in getattr(table, "merge_ranges", []) or []:
            parts = range_str.split(":")
            if len(parts) != 2:
                continue
            r0, c0 = _a1_to_rowcol(parts[0])
            r1, c1 = _a1_to_rowcol(parts[1])
            if r0 is None or r1 is None:
                continue
            if r0 <= row_idx <= r1 and c0 <= col_idx <= c1:
                origin = table.cell(r0, c0)
                if origin.value is not None:
                    return (
                        str(origin.formatted_value)
                        if getattr(origin, "formatted_value", None) is not None
                        else str(origin.value)
                    )
                return ""
    except Exception:
        pass
    return ""


def _numbers_table_to_matrix(table) -> list:
    """将 Numbers 的一个 Table 转为二维矩阵。"""
    try:
        raw_rows = table.rows(values_only=False)
    except Exception:
        raw_rows = []
    if not raw_rows:
        return []
    matrix = []
    for ri, row_cells in enumerate(raw_rows):
        row_data = []
        for ci, cell in enumerate(row_cells):
            val = _numbers_cell_display_value(cell, table, ri, ci)
            row_data.append(val.strip() if isinstance(val, str) else str(val))
        matrix.append(row_data)
    # 统一列数
    max_cols = max(len(r) for r in matrix) if matrix else 0
    for row in matrix:
        while len(row) < max_cols:
            row.append("")
    return matrix


def _load_numbers_sheets(path: Path):
    """从 Numbers 文件加载所有 sheet，逐个 yield (sheet_name, [matrix, ...])。"""
    doc = NumbersDocument(str(path))
    for sheet in doc.sheets:
        sheet_name = sheet.name
        matrices = []
        for table in sheet.tables:
            matrix = _numbers_table_to_matrix(table)
            if matrix:
                matrices.append(matrix)
        if not matrices:
            matrices = [[]]  # 空 sheet 也保留一节
        yield sheet_name, matrices


# ---------- Word 输出：三级大纲 + 大纲符 + 格式美化 ----------

def _set_paragraph_style(p, font_name="宋体", font_size_pt=10.5, left_indent_cm=0, space_after_pt=6):
    """统一设置段落字体与缩进。"""
    p.paragraph_format.left_indent = Cm(left_indent_cm)
    p.paragraph_format.space_after = Pt(space_after_pt)
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        try:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_name)
        except Exception:
            pass


def _add_heading_with_style(doc, text: str, level: int, symbol: str) -> None:
    """添加带大纲符和格式的标题（一级/二级/三级）。"""
    full_text = symbol + text
    p = doc.add_heading(full_text, level=level)
    for run in p.runs:
        run.font.name = "宋体"
        run.font.bold = True
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        try:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), "宋体")
        except Exception:
            pass
    p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.5 if level == 2 else 1.0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5 if level == 2 else 3)


def _is_likely_header_row(row: list, max_header_len: int = 50) -> bool:
    """判断一行是否像表头：非空、较短、无纯数字。"""
    if not row:
        return False
    for cell in row:
        s = (cell or "").strip()
        if not s:
            continue
        if len(s) > max_header_len:
            return False
        if s.replace(".", "").replace("-", "").isdigit():
            return False
    return True


def _matrix_to_outline_text(doc, matrix: list, l2_row_start: int) -> int:
    """
    将一张矩阵转为三级大纲写入 Word。
    每一行：第一列 → 二级大纲（带（一）（二））；其余列 → 三级内容（带①②③），列名：内容。
    返回下一行的二级序号（多表连续编号）。
    """
    if not matrix:
        return l2_row_start
    headers = None
    data_start = 0
    if len(matrix) > 1 and _is_likely_header_row(matrix[0]):
        headers = [str(c).strip() or "" for c in matrix[0]]
        data_start = 1
    ncols = len(matrix[0]) if matrix else 0
    l2_idx = l2_row_start
    for r in range(data_start, len(matrix)):
        row = matrix[r]
        if not row:
            continue
        str_row = []
        for j in range(max(len(row), ncols)):
            c = row[j] if j < len(row) else ""
            str_row.append((str(c).strip() if c is not None else ""))
        first_cell = str_row[0] if str_row else ""
        rest_cells = [(j, str_row[j]) for j in range(1, len(str_row)) if j < len(str_row) and str_row[j]]
        if first_cell:
            _add_heading_with_style(doc, first_cell, level=2, symbol=_outline_symbol_l2(l2_idx))
            l2_idx += 1
            for k, (j, val) in enumerate(rest_cells):
                label = (headers[j] + "：" if headers and j < len(headers) and headers[j] else "")
                third_text = _outline_symbol_l3(k) + (label + val if label else val)
                p = doc.add_paragraph(third_text.replace("\x00", ""))
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    try:
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn("w:eastAsia"), "宋体")
                    except Exception:
                        pass
            doc.add_paragraph()
        else:
            for j, val in rest_cells:
                label = (headers[j] + "：") if headers and j < len(headers) and headers[j] else ""
                p = doc.add_paragraph((label + val).replace("\x00", ""))
                p.paragraph_format.left_indent = Cm(1.0)
                _set_paragraph_style(p, left_indent_cm=1.0)
            if rest_cells:
                doc.add_paragraph()
    return l2_idx


def _add_section_to_doc(doc, section_name: str, matrices: list, l1_idx: int) -> None:
    """写入一节：一级大纲（带一、二、三符）+ 其下所有矩阵按三级大纲输出。"""
    _add_heading_with_style(doc, section_name, level=1, symbol=_outline_symbol_l1(l1_idx))
    if not matrices or all(not m for m in matrices):
        p = doc.add_paragraph("（本页无数据）")
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()
        return
    l2_start = 0
    for matrix in matrices:
        l2_start = _matrix_to_outline_text(doc, matrix, l2_start)


def convert_to_word(input_path: str, word_path: str = None) -> str:
    """
    将 Excel (.xlsx) 或 Numbers (.numbers) 工作簿转为一份 Word 文档。

    :param input_path: 输入文件路径（.xlsx 或 .numbers）
    :param word_path: 输出 Word 路径，默认与输入同目录、同名 .docx
    :return: 生成的 Word 文件路径
    """
    input_path = Path(input_path).resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_path}")

    if word_path is None:
        word_path = input_path.with_suffix(".docx")
    else:
        word_path = Path(word_path).resolve()

    suffix = input_path.suffix.lower()
    if suffix == ".xlsx":
        if openpyxl is None:
            raise RuntimeError("请先安装 openpyxl: pip install openpyxl")
        sheets_data = _load_excel_sheets(input_path)
    elif suffix == ".numbers":
        if NumbersDocument is None:
            raise RuntimeError("请先安装 numbers-parser: pip install numbers-parser")
        sheets_data = _load_numbers_sheets(input_path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .xlsx 与 .numbers")

    if WordDocument is None:
        raise RuntimeError("请先安装 python-docx: pip install python-docx")

    doc = WordDocument()
    # 正文与各级标题默认字体
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            s = doc.styles[style_name]
            s.font.name = "宋体"
            s.font.size = Pt(16 if style_name == "Heading 1" else 14 if style_name == "Heading 2" else 12 if style_name == "Heading 3" else 10.5)
            try:
                rPr = s._element.rPr
                if rPr is not None and hasattr(rPr, "get_or_add_rFonts"):
                    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass
        except Exception:
            pass

    for l1_idx, (sheet_name, matrices) in enumerate(sheets_data):
        _add_section_to_doc(doc, sheet_name, matrices, l1_idx)

    doc.save(str(word_path))
    return str(word_path)


def convert_excel_to_word(excel_path: str, word_path: str = None) -> str:
    """
    将 Excel 工作簿转为 Word 文档（兼容旧接口）。
    推荐使用 convert_to_word() 以同时支持 .xlsx 与 .numbers。
    """
    return convert_to_word(excel_path, word_path)


def main():
    parser = argparse.ArgumentParser(
        description="将 Excel (.xlsx) 或 Numbers (.numbers) 工作簿全部 sheet 转为一份 Word 产品技术手册（每 sheet 对应一级大纲）"
    )
    parser.add_argument(
        "input_file",
        help="输入的表格文件路径：.xlsx（Excel）或 .numbers（Apple Numbers）"
    )
    parser.add_argument(
        "-o", "--output",
        default=None,
        help="输出的 Word 文件路径；不指定则与输入文件同目录、同名 .docx"
    )
    args = parser.parse_args()

    try:
        out = convert_to_word(args.input_file, args.output)
        print(f"已生成 Word 文档: {out}")
    except FileNotFoundError as e:
        print(e, file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(e, file=sys.stderr)
        sys.exit(1)
    except RuntimeError as e:
        print(e, file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"转换失败: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
